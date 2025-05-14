# modules/output/final_output_generator.py

import os
import re
import json
from typing import List, Dict, Any, Optional, Tuple, Union, IO, BinaryIO, cast
from pathlib import Path
from modules.utils.logger import CustomLogger

# Try to import docx
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    # We don't need a fallback Document function since we'll check DOCX_AVAILABLE before using it

class FinalOutputGenerator:
    """
    Generate a comprehensive final output document combining the original
    modern play structure with translated Shakespeare lines.
    """
    
    def __init__(self, logger: Optional[CustomLogger] = None):
        self.logger = logger or CustomLogger("FinalOutputGenerator")
        self.logger.info("Initializing FinalOutputGenerator")
        
        if not DOCX_AVAILABLE:
            self.logger.warning("python-docx is not installed. Install with: pip install python-docx")
        
        # Regular expressions for parsing markdown play files
        self.act_pattern = re.compile(r'^ACT\s+([IVX\d]+)', re.IGNORECASE)
        self.scene_pattern = re.compile(r'^SCENE\s+([IVX\d]+)', re.IGNORECASE)
        self.stage_dir_pattern = re.compile(r'^\[(.*?)\]$')
        self.character_pattern = re.compile(r'^([A-Z][A-Z\s\-\']+)$')
        
    def generate_final_document(
        self, 
        modern_play_path: str,
        translations_dir: str,
        output_path: str = "outputs/final_output.docx",
        specific_act: Optional[str] = None,
        specific_scene: Optional[str] = None
    ) -> str:
        """
        Generate a comprehensive final document with translations and modern text.
        
        Args:
            modern_play_path: Path to the original modern play markdown file
            translations_dir: Directory containing the translated scene JSON files
            output_path: Path where the final Word document will be saved
            specific_act: If provided, only process this act
            specific_scene: If provided, only process this scene (requires specific_act)
            
        Returns:
            Path to the created document
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for generating Word documents. Install with: pip install python-docx")
            
        self.logger.info(f"Generating final document from {modern_play_path}")
        
        # If specific act/scene are provided, log the scope
        scope_msg = "full play"
        if specific_act:
            scope_msg = f"Act {specific_act}"
            if specific_scene:
                scope_msg += f", Scene {specific_scene}"
        self.logger.info(f"Processing scope: {scope_msg}")
        
        # Load the translations
        translations = self._load_translations(
            translations_dir, 
            specific_act=specific_act, 
            specific_scene=specific_scene
        )
        self.logger.info(f"Loaded translations for {len(translations)} scenes")
        
        # Create the document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = "Shakespeare Translation" 
        doc.core_properties.author = "AI Translation System"
        
        # Set up styles
        self._setup_document_styles(doc)
        
        # Add title based on scope
        title_text = "Shakespeare Translation"
        if specific_act:
            title_text += f" - Act {specific_act}"
            if specific_scene:
                title_text += f", Scene {specific_scene}"
                
        title = doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process the modern play and interleave translations
        self._process_play_file(
            doc, 
            modern_play_path, 
            translations,
            specific_act=specific_act,
            specific_scene=specific_scene
        )
        
        # Save the document
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        self.logger.info(f"Final document saved to: {output_path}")
        return output_path
        
    def _load_translations(
        self, 
        translations_dir: str,
        specific_act: Optional[str] = None,
        specific_scene: Optional[str] = None
    ) -> Dict[Tuple[str, str], List[Dict[str, Any]]]:
        """
        Load translated scenes from JSON files in the specified directory.
        
        Args:
            translations_dir: Directory containing translation JSON files
            specific_act: If provided, only load translations for this act
            specific_scene: If provided, only load translations for this scene
            
        Returns:
            A dictionary mapping (act, scene) tuples to lists of translated lines.
        """
        translations = {}
        
        # Handle the case of a specific translation file
        if specific_act and specific_scene:
            specific_file = f"act_{specific_act.lower()}_scene_{specific_scene.lower()}.json"
            filepath = Path(translations_dir) / specific_file
            
            if filepath.exists():
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    
                    # Extract the translated lines
                    if "translated_lines" in data:
                        translations[(specific_act.lower(), specific_scene.lower())] = data["translated_lines"]
                        self.logger.debug(f"Loaded {len(data['translated_lines'])} translated lines for Act {specific_act}, Scene {specific_scene}")
                except Exception as e:
                    self.logger.error(f"Error loading specific translation file {filepath}: {e}")
            else:
                self.logger.warning(f"Specific translation file not found: {filepath}")
            
            return translations
            
        # Load all relevant translation files
        for filepath in Path(translations_dir).glob("*.json"):
            match = re.match(r'act_([^_]+)_scene_(\w+)\.json', filepath.name)
            if not match:
                self.logger.warning(f"Skipping file with non-matching pattern: {filepath.name}")
                continue
                
            act, scene = match.groups()
            
            # Filter by specific act if provided
            if specific_act and act.lower() != specific_act.lower():
                continue
                
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    
                # Extract the translated lines
                if "translated_lines" in data:
                    translations[(act, scene)] = data["translated_lines"]
                    self.logger.debug(f"Loaded {len(data['translated_lines'])} translated lines for Act {act}, Scene {scene}")
            except Exception as e:
                self.logger.error(f"Error loading translation file {filepath}: {e}")
                
        return translations
        
    def _setup_document_styles(self, doc) -> None:
        """Set up styles for the document."""
        # Act style
        act_style = doc.styles.add_style('Act', 1)
        act_style.font.bold = True
        act_style.font.size = Pt(16)
        act_style.paragraph_format.space_before = Pt(24)
        act_style.paragraph_format.space_after = Pt(12)
        
        # Scene style
        scene_style = doc.styles.add_style('Scene', 1)
        scene_style.font.bold = True
        scene_style.font.size = Pt(14)
        scene_style.paragraph_format.space_before = Pt(18)
        scene_style.paragraph_format.space_after = Pt(12)
        
        # Stage direction style
        stage_dir_style = doc.styles.add_style('StageDirection', 1)
        stage_dir_style.font.italic = True
        stage_dir_style.paragraph_format.space_before = Pt(6)
        stage_dir_style.paragraph_format.space_after = Pt(6)
        
        # Character style
        character_style = doc.styles.add_style('Character', 1)
        character_style.font.bold = True
        character_style.paragraph_format.space_before = Pt(12)
        character_style.paragraph_format.space_after = Pt(0)
        
        # Dialogue style
        dialogue_style = doc.styles.add_style('Dialogue', 1)
        dialogue_style.paragraph_format.space_before = Pt(0)
        dialogue_style.paragraph_format.space_after = Pt(12)
        
    def _process_play_file(
        self, 
        doc, 
        filepath: str, 
        translations: Dict[Tuple[str, str], List[Dict[str, Any]]],
        specific_act: Optional[str] = None,
        specific_scene: Optional[str] = None
    ) -> None:
        """
        Process the modern play file and interleave it with translations.
        
        Uses fuzzy matching between modern lines in the play file and modern lines in the translations.
        """
        self.logger.info(f"Processing modern play file: {filepath}")
        
        # Function for fuzzy text matching
        def fuzzy_match(text1, text2, threshold=0.75):
            # Normalize texts for comparison
            def normalize(text):
                if text is None:
                    return ""
                # Convert to lowercase
                text = text.lower()
                # Remove punctuation
                text = ''.join(c for c in text if c.isalnum() or c.isspace())
                # Normalize whitespace
                text = ' '.join(text.split())
                return text
                
            norm1 = normalize(text1)
            norm2 = normalize(text2)
            
            # For very short texts, require exact match after normalization
            if len(norm1) < 10 or len(norm2) < 10:
                return norm1 == norm2
                
            # Simple similarity ratio
            # Count shared words
            words1 = set(norm1.split())
            words2 = set(norm2.split())
            common_words = words1.intersection(words2)
            
            # Calculate Jaccard similarity
            if len(words1) + len(words2) == 0:
                return False
            similarity = len(common_words) / (len(words1) + len(words2) - len(common_words))
            
            return similarity >= threshold
        
        # Function to detect markdown headers and extract act/scene info
        def parse_markdown_header(line):
            # Check for markdown act header (# ACT X)
            if line.startswith("# ") and "ACT" in line.upper():
                parts = line.upper().replace("#", "").strip().split()
                if len(parts) >= 2 and parts[0] == "ACT":
                    return "act", parts[1]
            
            # Check for markdown scene header (## SCENE X)
            if line.startswith("## ") and "SCENE" in line.upper():
                parts = line.upper().replace("#", "").strip().split()
                if len(parts) >= 2 and parts[0] == "SCENE":
                    return "scene", parts[1]
                    
            return None, None
        
        # Function to detect character names more robustly, excluding markdown headers
        def is_character_line(line):
            # Skip markdown headers
            if line.startswith("#"):
                return False
                
            # If it's in all caps and not too long, it's likely a character name
            if line.isupper() and len(line.split()) <= 4:
                return True
                
            # Check against the character pattern
            if self.character_pattern.match(line):
                return True
                
            # Additional checks (e.g., names might be followed by a colon)
            if ":" in line and line.split(":")[0].strip().isupper():
                return True
                
            return False
        
        # Track current context
        current_act = None
        current_scene = None
        current_character = None
        
        # Track whether we're in the target section for specific processing
        in_target_section = specific_act is None  # Initially true if processing everything
        
        # Keep track of translation indices for each scene
        translation_indices = {}  # (act, scene) -> current index
        
        # Read the file
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                lines = f.readlines()
        except Exception as e:
            self.logger.error(f"Error reading modern play file: {e}")
            return
        
        # Process each line
        for line_num, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # IMPROVED HEADER DETECTION FOR MARKDOWN
            header_type, header_value = parse_markdown_header(line)
            
            # Handle markdown act header
            if header_type == "act":
                current_act = header_value
                current_scene = None
                
                self.logger.info(f"Detected markdown Act header: {line} -> Act {current_act}")
                
                # Check if this is the specific act we're looking for
                if specific_act:
                    # Add null check before using lower()
                    act_lower = current_act.lower() if current_act is not None else ""
                    specific_act_lower = specific_act.lower() if specific_act is not None else ""
                    in_target_section = act_lower == specific_act_lower
                    if not in_target_section:
                        self.logger.debug(f"Skipping Act {current_act} (not in target)")
                        continue
                
                # Add act header to document
                if in_target_section:
                    # Use string conversion to handle potential None values
                    act_para = doc.add_paragraph("ACT " + str(current_act), style='Act')
                    self.logger.debug(f"Added Act {current_act}")
                
                continue

            # For scene headers
            if header_type == "scene":
                current_scene = header_value
                
                self.logger.info(f"Detected markdown Scene header: {line} -> Scene {current_scene}")
                
                # Reset translation index for new scene
                act_key = current_act.lower() if current_act is not None else ""
                scene_key = current_scene.lower() if current_scene is not None else ""
                translation_indices[(act_key, scene_key)] = 0
                
                # Check if this is the specific scene we're looking for
                if specific_scene and specific_act:
                    # Add null check before using lower()
                    scene_lower = current_scene.lower() if current_scene is not None else ""
                    specific_scene_lower = specific_scene.lower() if specific_scene is not None else ""
                    in_target_section = scene_lower == specific_scene_lower
                    if not in_target_section:
                        self.logger.debug(f"Skipping Scene {current_scene} (not in target)")
                        continue
                
                # Add scene header to document if we're processing this section
                if in_target_section:
                    # Use string conversion to handle potential None values
                    scene_para = doc.add_paragraph("SCENE " + str(current_scene), style='Scene')
                    self.logger.debug(f"Added Scene {current_scene}")
                
                continue

            # LEGACY ACT/SCENE DETECTION (for non-markdown format)
            act_match = self.act_pattern.match(line) if hasattr(self, 'act_pattern') else None
            if act_match:
                current_act = act_match.group(1)
                current_scene = None
                
                # Check if this is the specific act we're looking for
                if specific_act:
                    in_target_section = current_act.lower() == specific_act.lower()
                    if not in_target_section:
                        self.logger.debug(f"Skipping Act {current_act} (not in target)")
                        continue
                
                # Add act header to document
                if in_target_section:
                    act_para = doc.add_paragraph("ACT " + current_act, style='Act')
                    self.logger.debug(f"Added Act {current_act}")
                
                continue
                
            # Skip if we're not in the target act
            if specific_act and not in_target_section:
                continue
                
            scene_match = self.scene_pattern.match(line) if hasattr(self, 'scene_pattern') else None
            if scene_match:
                current_scene = scene_match.group(1)
                
                # Reset translation index for new scene
                act_key = current_act.lower() if current_act is not None else ""
                scene_key = current_scene.lower() if current_scene is not None else ""
                translation_indices[(act_key, scene_key)] = 0
                
                # Check if this is the specific scene we're looking for
                if specific_scene and specific_act:
                    # Only process the specific scene in the specific act
                    in_target_section = (current_scene.lower() == specific_scene.lower())
                    if not in_target_section:
                        self.logger.debug(f"Skipping Scene {current_scene} (not in target)")
                        continue
                
                # Add scene header to document if we're processing this section
                if in_target_section:
                    scene_para = doc.add_paragraph("SCENE " + current_scene, style='Scene')
                    self.logger.debug(f"Added Scene {current_scene}")
                
                continue
                
            # Skip if we're not in the target scene
            if specific_scene and specific_act and not in_target_section:
                continue
                
            # Process structural and dialogue lines if we're in the target section
            if in_target_section:
                # Check if it's a stage direction
                stage_dir_match = self.stage_dir_pattern.match(line)
                if stage_dir_match:
                    direction = stage_dir_match.group(1)
                    
                    # Add stage direction to document
                    direction_para = doc.add_paragraph(f"[{direction}]", style='StageDirection')
                    self.logger.debug(f"Added stage direction: [{direction}]")
                    continue
                    
                # IMPROVED CHARACTER DETECTION
                if is_character_line(line):
                    # Extract character name (remove any colon)
                    if ":" in line:
                        current_character = line.split(":")[0].strip()
                    else:
                        current_character = line.strip()
                    
                    # Add character name to document
                    char_para = doc.add_paragraph(current_character, style='Character')
                    self.logger.debug(f"Added character: {current_character}")
                    continue
                    
                # If it's none of the above, it's a dialogue line
                if current_character and current_act is not None and current_scene is not None:
                    # Get normalized act and scene keys
                    act_key = current_act.lower() if current_act is not None else ""
                    scene_key = current_scene.lower() if current_scene is not None else ""
                    
                    self.logger.debug(f"Processing dialogue in act_{act_key}_scene_{scene_key}")
                    
                    # Find matching translation by modern content
                    translation = None
                    search_window = 8  # Increased window size for better matching
                    
                    # Look up translations for this act and scene
                    scene_translations = translations.get((act_key, scene_key), [])
                    current_index = translation_indices.get((act_key, scene_key), 0)
                    
                    # Log some stats
                    self.logger.debug(f"Act {current_act}, Scene {current_scene}: Searching for match at index {current_index}/{len(scene_translations)}")
                    
                    if scene_translations and current_index < len(scene_translations):
                        # Look for an exact or fuzzy match in the next few translations
                        search_end = min(current_index + search_window, len(scene_translations))
                        
                        for i in range(current_index, search_end):
                            t = scene_translations[i]
                            if "original_modern_line" in t:
                                modern_line = t["original_modern_line"]
                                
                                # Try fuzzy matching
                                if fuzzy_match(modern_line, line):
                                    translation = t
                                    # Update the index for next time
                                    translation_indices[(act_key, scene_key)] = i + 1
                                    self.logger.debug(f"Found fuzzy match at index {i}: '{modern_line[:30]}...' ≈ '{line[:30]}...'")
                                    break
                    
                    # If no match found in window, try a broader search as fallback
                    if translation is None and scene_translations:
                        self.logger.debug(f"No match in window, trying broader search for: '{line[:30]}...'")
                        
                        # Look through all remaining translations
                        for i in range(current_index, len(scene_translations)):
                            t = scene_translations[i]
                            if "original_modern_line" in t:
                                modern_line = t["original_modern_line"]
                                
                                # Try fuzzy matching with lower threshold for the last act
                                threshold = 0.6 if act_key.lower() == "v" else 0.75
                                if fuzzy_match(modern_line, line, threshold):
                                    translation = t
                                    # Update the index for next time
                                    translation_indices[(act_key, scene_key)] = i + 1
                                    self.logger.debug(f"Found fuzzy match at index {i}: '{modern_line[:30]}...' ≈ '{line[:30]}...'")
                                    break
                    
                    # Create a table for this line (Shakespeare | References | Modern)
                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Table Grid'
                    
                    # Set column widths if there are columns to set widths for
                    if hasattr(table, 'columns') and len(table.columns) >= 3:
                        for i, width in enumerate([2.5, 2.0, 2.5]):
                            # Make sure we have cells in this column before iterating
                            col = table.columns[i]
                            if hasattr(col, 'cells') and len(col.cells) > 0:
                                for cell in col.cells:
                                    cell.width = Inches(width)
                    
                    # Make sure the table has rows before trying to access them
                    if hasattr(table, 'rows') and len(table.rows) > 0:
                        # Get first row cells
                        row = table.rows[0]
                        if hasattr(row, 'cells') and len(row.cells) >= 3:
                            cells = row.cells
                            
                            # Column 1: Shakespeare text or placeholder
                            if translation and "text" in translation:
                                cells[0].text = translation["text"]
                            else:
                                cells[0].text = "[No translation available]"
                                self.logger.warning(f"No translation match found for: '{line[:50]}...'")
                                
                            # Column 2: References
                            if translation and "formatted_references" in translation:
                                cells[1].text = "\n".join(translation["formatted_references"])
                            elif translation and "references" in translation:
                                # Format the references if they haven't been pre-formatted
                                formatted_refs = []
                                for ref in translation["references"]:
                                    title = ref.get("title", "Unknown")
                                    act_ref = ref.get("act", "")
                                    scene_ref = ref.get("scene", "")
                                    line_ref = ref.get("line", "")
                                    formatted_refs.append(f"{title} ({act_ref}.{scene_ref}.{line_ref})")
                                cells[1].text = "\n".join(formatted_refs)
                            else:
                                cells[1].text = ""
                                
                            # Column 3: Modern text
                            cells[2].text = line
                            
                            # Format the table
                            for cell in row.cells:
                                if hasattr(cell, 'paragraphs'):
                                    for paragraph in cell.paragraphs:
                                        paragraph.style = 'Dialogue'
                    
                    self.logger.debug(f"Added dialogue line table for: {line[:30]}...")
        
        self.logger.info("Completed processing play file")

    def generate_scene_document(
        self, 
        act: str, 
        scene: str,
        modern_play_path: str, 
        translation_file: str, 
        output_path: Optional[str] = None
    ) -> str:
        """
        Generate a document for a specific scene from a specific translation file.
        
        Args:
            act: Act identifier
            scene: Scene identifier
            modern_play_path: Path to the modern play markdown file
            translation_file: Path to the specific translation JSON file
            output_path: Path for the output document (default: auto-generated)
            
        Returns:
            Path to the created document
        """
        if not output_path:
            output_path = f"outputs/translated_act_{act}_scene_{scene}.docx"
            
        # Load the specific translation file
        translations_dir = os.path.dirname(translation_file)
        
        return self.generate_final_document(
            modern_play_path=modern_play_path,
            translations_dir=translations_dir,
            output_path=output_path,
            specific_act=act,
            specific_scene=scene
        )
    
    def generate_from_line_list(
        self, 
        act: str,
        scene: str,
        modern_lines: List[str],
        translated_lines: List[Dict[str, Any]],
        output_path: str = "outputs/translated_excerpt.docx"
    ) -> str:
        """
        Generate a document from explicit lists of modern and translated lines.
        Useful for UI operations where you might not have a full scene or play file.
        
        Args:
            act: Act identifier
            scene: Scene identifier
            modern_lines: List of modern lines
            translated_lines: List of translation data dictionaries
            output_path: Path for the output document
            
        Returns:
            Path to the created document
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for generating Word documents. Install with: pip install python-docx")
            
        # Create the document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = f"Shakespeare Translation - Act {act}, Scene {scene} (Excerpt)"
        doc.core_properties.author = "AI Translation System"
        
        # Set up styles
        self._setup_document_styles(doc)
        
        # Add title
        title = doc.add_heading(f"Shakespeare Translation - Act {act}, Scene {scene} (Excerpt)", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add act and scene headers
        doc.add_paragraph(f"ACT {act}", style='Act')
        doc.add_paragraph(f"SCENE {scene}", style='Scene')
        
        # Process the lines
        for i, modern_line in enumerate(modern_lines):
            # Get the corresponding translation if available
            translation = None
            if i < len(translated_lines):
                translation = translated_lines[i]
            
            # Create a table for this line
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Set column widths if we can
            if hasattr(table, 'columns') and len(table.columns) >= 3:
                for i, width in enumerate([2.5, 2.0, 2.5]):
                    col = table.columns[i]
                    if hasattr(col, 'cells') and len(col.cells) > 0:
                        for cell in col.cells:
                            cell.width = Inches(width)
            
            # Make sure the table has rows before trying to access them
            if hasattr(table, 'rows') and len(table.rows) > 0:
                row = table.rows[0]
                if hasattr(row, 'cells') and len(row.cells) >= 3:
                    cells = row.cells
                    
                    # Column 1: Shakespeare text or placeholder
                    if translation and "text" in translation:
                        cells[0].text = translation["text"]
                    else:
                        cells[0].text = "[No translation available]"
                        
                    # Column 2: References
                    if translation and "formatted_references" in translation:
                        cells[1].text = "\n".join(translation["formatted_references"])
                    elif translation and "references" in translation:
                        # Format the references if they haven't been pre-formatted
                        formatted_refs = []
                        for ref in translation["references"]:
                            title = ref.get("title", "Unknown")
                            act_ref = ref.get("act", "")
                            scene_ref = ref.get("scene", "")
                            line_ref = ref.get("line", "")
                            formatted_refs.append(f"{title} ({act_ref}.{scene_ref}.{line_ref})")
                        cells[1].text = "\n".join(formatted_refs)
                    else:
                        cells[1].text = ""
                        
                    # Column 3: Modern text
                    cells[2].text = modern_line
                    
                    # Format the table
                    for cell in row.cells:
                        if hasattr(cell, 'paragraphs'):
                            for paragraph in cell.paragraphs:
                                paragraph.style = 'Dialogue'
        
        # Save the document
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        self.logger.info(f"Excerpt document saved to: {output_path}")
        return output_path