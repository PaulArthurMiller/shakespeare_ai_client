# format_translated_play.py

import os
import json
import argparse
from pathlib import Path
from typing import List, Dict, Any, Optional, Union, cast, TextIO, BinaryIO

import re

class PlayFormatter:
    """
    A dedicated formatter to convert the JSON output of translated scenes
    into various presentation formats.
    """
    
    def __init__(self, json_dir: str, output_dir: str):
        """
        Initialize the formatter.
        
        Args:
            json_dir: Directory containing the JSON scene files
            output_dir: Directory where formatted outputs will be saved
        """
        self.json_dir = Path(json_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Regular expression to extract act and scene numbers from filenames
        self.filename_pattern = re.compile(r'act_([^_]+)_scene_(\w+)\.json')
        
    def _load_scene_files(self) -> List[Dict[str, Any]]:
        """
        Load all scene JSON files and sort them by act and scene number.
        
        Returns:
            A list of tuples with (act, scene, scene_data)
        """
        scene_files = []
        
        # List all JSON files in the directory
        for filepath in sorted(self.json_dir.glob('*.json')):
            match = self.filename_pattern.match(filepath.name)
            if not match:
                print(f"Skipping file with non-matching pattern: {filepath.name}")
                continue
                
            act, scene = match.groups()
            
            # Try to convert act and scene to numeric values for sorting
            try:
                # Handle Roman numerals
                if re.match(r'^[IVXLCDM]+$', act.upper()):
                    act_num = self._roman_to_int(act.upper())
                else:
                    act_num = float(act)
                    
                if re.match(r'^[IVXLCDM]+$', scene.upper()):
                    scene_num = self._roman_to_int(scene.upper())
                else:
                    scene_num = float(scene)
            except ValueError:
                # If conversion fails, use string values with high sort priority
                act_num = act
                scene_num = scene
            
            # Load the scene data
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    scene_data = json.load(f)
                    # Ensure scene_data is a dictionary
                    if isinstance(scene_data, dict):
                        scene_files.append((act_num, scene_num, act, scene, scene_data))
                    else:
                        print(f"Warning: {filepath} does not contain a valid JSON object")
            except Exception as e:
                print(f"Error loading {filepath}: {e}")
        
        # Sort by act and scene
        scene_files.sort()
        return scene_files
    
    def _roman_to_int(self, roman: str) -> int:
        """Convert Roman numeral to integer."""
        values = {
            'I': 1, 'V': 5, 'X': 10, 'L': 50, 
            'C': 100, 'D': 500, 'M': 1000
        }
        total = 0
        prev = 0
        
        for char in reversed(roman):
            current = values[char]
            if current >= prev:
                total += current
            else:
                total -= current
            prev = current
            
        return total
    
    def format_markdown(self, output_filename: str = "translated_play.md") -> str:
        """
        Format all scenes as a markdown document with two-column format.
        
        Args:
            output_filename: Name of the output markdown file
            
        Returns:
            Path to the created file
        """
        output_path = self.output_dir / output_filename
        
        # Load and sort scene files
        scene_files = self._load_scene_files()
        
        with open(output_path, 'w', encoding='utf-8') as outfile:
            outfile.write("# The Translated Play\n\n")
            
            for _, _, act, scene, scene_data in scene_files:
                # Make sure scene_data is a dictionary
                if not isinstance(scene_data, dict):
                    continue
                
                # Write act and scene headers
                outfile.write(f"## ACT {act.upper()}\n\n")
                outfile.write(f"### SCENE {scene.upper()}\n\n")
                
                # Create table header
                outfile.write("| Shakespearean Text | Source References | Modern Text |\n")
                outfile.write("|-------------------|-------------------|------------|\n")
                
                # Write each line
                translated_lines = scene_data.get("translated_lines", [])
                if isinstance(translated_lines, list):
                    for line in translated_lines:
                        if not isinstance(line, dict):
                            continue
                            
                        # Get the Shakespearean text
                        shakespeare_text = line.get("text", "").replace("\n", " ")
                        
                        # Get the references
                        formatted_refs = line.get("formatted_references", [])
                        if not formatted_refs and "references" in line:
                            # Format references if they haven't been pre-formatted
                            references = line.get("references", [])
                            if isinstance(references, list):
                                formatted_refs = []
                                for ref in references:
                                    if isinstance(ref, dict):
                                        source = ref.get("title", "Unknown")
                                        act_ref = ref.get("act", "")
                                        scene_ref = ref.get("scene", "")
                                        line_ref = ref.get("line", "")
                                        formatted_refs.append(f"{source} ({act_ref}.{scene_ref}.{line_ref})")
                        
                        if isinstance(formatted_refs, list):
                            refs_text = "<br>".join(formatted_refs)
                        else:
                            refs_text = ""
                        
                        # Get the original modern text
                        modern_text = line.get("original_modern_line", "").replace("\n", " ")
                        
                        # Escape pipe characters for markdown tables
                        shakespeare_text = shakespeare_text.replace("|", "\\|")
                        refs_text = refs_text.replace("|", "\\|")
                        modern_text = modern_text.replace("|", "\\|")
                        
                        # Write the table row
                        outfile.write(f"| {shakespeare_text} | {refs_text} | {modern_text} |\n")
                
                # Add a separator between scenes
                outfile.write("\n---\n\n")
        
        print(f"Markdown formatted play saved to: {output_path}")
        return str(output_path)
    
    def format_docx(self, output_filename: str = "translated_play.docx") -> str:
        """
        Format all scenes as a Word document.
        
        Args:
            output_filename: Name of the output Word file
            
        Returns:
            Path to the created file
        """
        try:
            from docx import Document  # type: ignore
            from docx.shared import Pt, Inches  # type: ignore
            from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        except ImportError:
            print("python-docx not installed. Install with: pip install python-docx")
            return ""
        
        output_path = self.output_dir / output_filename
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = "Translated Shakespeare Play"
        doc.core_properties.author = "AI Translation System"
        
        # Title page
        title = doc.add_heading("Translated Play", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Load and sort scene files
        scene_files = self._load_scene_files()
        
        for _, _, act, scene, scene_data in scene_files:
            if not isinstance(scene_data, dict):
                continue
                
            # Act heading
            doc.add_heading(f"ACT {act.upper()}", level=1)
            
            # Scene heading
            doc.add_heading(f"SCENE {scene.upper()}", level=2)
            
            # Create table
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Set column widths
            for cell in table.columns[0].cells:
                cell.width = Inches(2.5)
            for cell in table.columns[1].cells:
                cell.width = Inches(2.0)
            for cell in table.columns[2].cells:
                cell.width = Inches(2.5)
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "Shakespearean Text"
            header_cells[1].text = "Source References"
            header_cells[2].text = "Modern Text"
            
            # Add each line to the table
            translated_lines = scene_data.get("translated_lines", [])
            if isinstance(translated_lines, list):
                for line in translated_lines:
                    if not isinstance(line, dict):
                        continue
                        
                    row_cells = table.add_row().cells
                    
                    # Shakespearean text
                    row_cells[0].text = line.get("text", "")
                    
                    # References
                    refs = []
                    references = line.get("references", [])
                    if isinstance(references, list):
                        for ref in references:
                            if isinstance(ref, dict):
                                source = ref.get("title", "Unknown")
                                act_ref = ref.get("act", "")
                                scene_ref = ref.get("scene", "")
                                line_ref = ref.get("line", "")
                                refs.append(f"{source} ({act_ref}.{scene_ref}.{line_ref})")
                    
                    row_cells[1].text = "\n".join(refs)
                    
                    # Modern text
                    row_cells[2].text = line.get("original_modern_line", "")
            
            # Add page break between scenes
            doc.add_page_break()
        
        # Save the document - convert Path to string
        doc.save(str(output_path))
        print(f"Word document formatted play saved to: {output_path}")
        return str(output_path)
    
    def format_html(self, output_filename: str = "translated_play.html") -> str:
        """
        Format all scenes as an HTML document.
        
        Args:
            output_filename: Name of the output HTML file
            
        Returns:
            Path to the created file
        """
        output_path = self.output_dir / output_filename
        
        # Load and sort scene files
        scene_files = self._load_scene_files()
        
        with open(output_path, 'w', encoding='utf-8') as outfile:
            # HTML header
            outfile.write("""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Translated Shakespeare Play</title>
    <style>
        body { font-family: 'Garamond', serif; margin: 40px; line-height: 1.6; }
        h1 { text-align: center; margin-bottom: 40px; }
        h2 { color: #4a4a4a; margin-top: 30px; }
        h3 { color: #666; }
        table { width: 100%; border-collapse: collapse; margin-bottom: 30px; }
        th { background-color: #f2f2f2; padding: 10px; text-align: left; border-bottom: 2px solid #ddd; }
        td { padding: 10px; border-bottom: 1px solid #ddd; vertical-align: top; }
        .shakespeare { width: 40%; }
        .references { width: 20%; font-size: 0.9em; color: #666; }
        .modern { width: 40%; }
        .separator { margin: 40px 0; border-top: 1px dashed #ccc; }
    </style>
</head>
<body>
    <h1>Translated Shakespeare Play</h1>
""")
            
            for _, _, act, scene, scene_data in scene_files:
                if not isinstance(scene_data, dict):
                    continue
                    
                # Act and scene headers
                outfile.write(f'    <h2>ACT {act.upper()}</h2>\n')
                outfile.write(f'    <h3>SCENE {scene.upper()}</h3>\n')
                
                # Table
                outfile.write('    <table>\n')
                outfile.write('        <tr>\n')
                outfile.write('            <th class="shakespeare">Shakespearean Text</th>\n')
                outfile.write('            <th class="references">Source References</th>\n')
                outfile.write('            <th class="modern">Modern Text</th>\n')
                outfile.write('        </tr>\n')
                
                translated_lines = scene_data.get("translated_lines", [])
                if isinstance(translated_lines, list):
                    for line in translated_lines:
                        if not isinstance(line, dict):
                            continue
                            
                        outfile.write('        <tr>\n')
                        
                        # Shakespearean text
                        shakespeare_text = line.get("text", "").replace("<", "&lt;").replace(">", "&gt;")
                        outfile.write(f'            <td class="shakespeare">{shakespeare_text}</td>\n')
                        
                        # References
                        refs = []
                        references = line.get("references", [])
                        if isinstance(references, list):
                            for ref in references:
                                if isinstance(ref, dict):
                                    source = ref.get("title", "Unknown").replace("<", "&lt;").replace(">", "&gt;")
                                    act_ref = ref.get("act", "")
                                    scene_ref = ref.get("scene", "")
                                    line_ref = ref.get("line", "")
                                    refs.append(f"{source} ({act_ref}.{scene_ref}.{line_ref})")
                        
                        refs_html = "<br>".join(refs)
                        outfile.write(f'            <td class="references">{refs_html}</td>\n')
                        
                        # Modern text
                        modern_text = line.get("original_modern_line", "").replace("<", "&lt;").replace(">", "&gt;")
                        outfile.write(f'            <td class="modern">{modern_text}</td>\n')
                        
                        outfile.write('        </tr>\n')
                
                outfile.write('    </table>\n')
                outfile.write('    <div class="separator"></div>\n')
            
            # HTML footer
            outfile.write("""</body>
</html>""")
        
        print(f"HTML formatted play saved to: {output_path}")
        return str(output_path)

def main():
    """Main function to run the formatter from command line."""
    parser = argparse.ArgumentParser(description='Format translated Shakespeare scenes')
    parser.add_argument('--json-dir', type=str, required=True, 
                        help='Directory containing JSON scene files')
    parser.add_argument('--output-dir', type=str, default='formatted_output',
                        help='Directory for output files (default: formatted_output)')
    parser.add_argument('--format', type=str, choices=['markdown', 'docx', 'html', 'all'], 
                        default='all', help='Output format (default: all)')
    
    args = parser.parse_args()
    
    formatter = PlayFormatter(args.json_dir, args.output_dir)
    
    if args.format == 'markdown' or args.format == 'all':
        formatter.format_markdown()
    
    if args.format == 'docx' or args.format == 'all':
        formatter.format_docx()
    
    if args.format == 'html' or args.format == 'all':
        formatter.format_html()

if __name__ == "__main__":
    main()