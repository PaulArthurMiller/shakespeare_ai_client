# modules/output/save_modern_play.py

import os
import json
import re
from typing import List, Dict, Any, Optional, Union, Tuple
from pathlib import Path

"""
This module handles saving modern play scenes and full plays to various formats.
It supports exporting to both DOCX and Markdown formats.
"""

class SceneExporter:
    """Handles exporting individual scenes to various formats."""
    
    def __init__(self):
        """Initialize the SceneExporter."""
        # Check if python-docx is available
        try:
            import docx
            self.docx_available = True
        except ImportError:
            self.docx_available = False
            print("Warning: python-docx not available. Install with: pip install python-docx")
    
    def export_scene_from_json(self, json_path: str, output_path: str) -> str:
        """
        Export a scene from a JSON file to DOCX format.
        
        Args:
            json_path: Path to the JSON scene file
            output_path: Path for the output file
            
        Returns:
            Path to the created file
        """
        if not self.docx_available:
            raise ImportError("python-docx is required for DOCX export")
        
        # Load the scene data
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                scene_data = json.load(f)
        except Exception as e:
            raise ValueError(f"Error loading scene JSON: {str(e)}")
        
        # Extract scene content and metadata
        act = scene_data.get("act", "Unknown")
        scene = scene_data.get("scene", "Unknown")
        script = scene_data.get("script", "")
        
        # Create DOCX
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Add title
        title = doc.add_heading(f"Act {act}, Scene {scene}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process the script content - split by lines
        lines = script.split('\n')
        
        # Add content
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if it's a stage direction [...]
            if line.startswith('[') and line.endswith(']'):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
                continue
                
            # Check if it's a character name (all caps)
            if line.isupper() and len(line.split()) <= 3:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                p.space_after = Pt(0)  # No space after character name
                continue
                
            # Regular dialogue
            p = doc.add_paragraph(line)
            p.paragraph_format.left_indent = Pt(36)  # Indent dialogue
        
        # Save the document
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        return output_path
    
    def export_scene_from_markdown(self, md_path: str, output_path: str) -> str:
        """
        Export a scene from a Markdown file to DOCX format.
        
        Args:
            md_path: Path to the Markdown scene file
            output_path: Path for the output file
            
        Returns:
            Path to the created file
        """
        if not self.docx_available:
            raise ImportError("python-docx is required for DOCX export")
        
        # Load the markdown content
        try:
            with open(md_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
        except Exception as e:
            raise ValueError(f"Error loading markdown file: {str(e)}")
        
        # Extract act and scene from filename or content
        act = "Unknown"
        scene = "Unknown"
        
        # Try to extract from filename
        filename = os.path.basename(md_path)
        match = re.search(r'act_([^_]+)_scene_(\w+)', filename)
        if match:
            act, scene = match.groups()
        
        # Try to extract from content if not found in filename
        if act == "Unknown" or scene == "Unknown":
            # Look for ACT and SCENE headers in the content
            act_match = re.search(r'ACT\s+([IVX\d]+)', md_content, re.IGNORECASE)
            scene_match = re.search(r'SCENE\s+([IVX\d]+)', md_content, re.IGNORECASE)
            
            if act_match:
                act = act_match.group(1)
            if scene_match:
                scene = scene_match.group(1)
        
        # Create DOCX
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Add title
        title = doc.add_heading(f"Act {act}, Scene {scene}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process the markdown content - split by lines
        lines = md_content.split('\n')
        
        # Skip act and scene headers
        content_lines = []
        for line in lines:
            if re.match(r'^ACT\s+[IVX\d]+', line, re.IGNORECASE):
                continue
            if re.match(r'^SCENE\s+[IVX\d]+', line, re.IGNORECASE):
                continue
            content_lines.append(line)
        
        # Add content
        for line in content_lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if it's a stage direction [...]
            if line.startswith('[') and line.endswith(']'):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
                continue
                
            # Check if it's a character name (all caps)
            if line.isupper() and len(line.split()) <= 3:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                p.space_after = Pt(0)  # No space after character name
                continue
                
            # Regular dialogue
            p = doc.add_paragraph(line)
            p.paragraph_format.left_indent = Pt(36)  # Indent dialogue
        
        # Save the document
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        return output_path


class PlayExporter:
    """Handles exporting full plays to various formats."""
    
    def __init__(self):
        """Initialize the PlayExporter."""
        # Check if python-docx is available
        try:
            import docx
            self.docx_available = True
        except ImportError:
            self.docx_available = False
            print("Warning: python-docx not available. Install with: pip install python-docx")
        
        # Create a scene exporter for individual scenes
        self.scene_exporter = SceneExporter()
    
    def export_play_from_scenes(self, scene_paths: List[str], output_path: str, title: str = "Play") -> str:
        """
        Export a complete play from individual scene files to DOCX format.
        
        Args:
            scene_paths: List of paths to scene JSON or MD files
            output_path: Path for the output file
            title: Title of the play
            
        Returns:
            Path to the created file
        """
        if not self.docx_available:
            raise ImportError("python-docx is required for DOCX export")
        
        # Create DOCX
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Add title page
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add break for title page
        doc.add_page_break()
        
        # Process each scene
        for scene_path in scene_paths:
            # Check if it's JSON or MD
            is_json = scene_path.lower().endswith('.json')
            
            try:
                if is_json:
                    # Load scene data
                    with open(scene_path, 'r', encoding='utf-8') as f:
                        scene_data = json.load(f)
                    
                    # Extract metadata
                    act = scene_data.get("act", "Unknown")
                    scene = scene_data.get("scene", "Unknown")
                    script = scene_data.get("script", "")
                    
                    # Add scene header
                    doc.add_heading(f"Act {act}, Scene {scene}", level=1)
                    
                    # Process script content
                    lines = script.split('\n')
                    
                    # Add content
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                            
                        # Check if it's a stage direction [...]
                        if line.startswith('[') and line.endswith(']'):
                            p = doc.add_paragraph()
                            run = p.add_run(line)
                            run.italic = True
                            continue
                            
                        # Check if it's a character name (all caps)
                        if line.isupper() and len(line.split()) <= 3:
                            p = doc.add_paragraph()
                            run = p.add_run(line)
                            run.bold = True
                            p.space_after = Pt(0)  # No space after character name
                            continue
                            
                        # Regular dialogue
                        p = doc.add_paragraph(line)
                        p.paragraph_format.left_indent = Pt(36)  # Indent dialogue
                
                else:
                    # Load markdown content
                    with open(scene_path, 'r', encoding='utf-8') as f:
                        md_content = f.read()
                    
                    # Extract act and scene from filename or content
                    act = "Unknown"
                    scene = "Unknown"
                    
                    # Try to extract from filename
                    filename = os.path.basename(scene_path)
                    match = re.search(r'act_([^_]+)_scene_(\w+)', filename)
                    if match:
                        act, scene = match.groups()
                    
                    # Try to extract from content if not found in filename
                    if act == "Unknown" or scene == "Unknown":
                        act_match = re.search(r'ACT\s+([IVX\d]+)', md_content, re.IGNORECASE)
                        scene_match = re.search(r'SCENE\s+([IVX\d]+)', md_content, re.IGNORECASE)
                        
                        if act_match:
                            act = act_match.group(1)
                        if scene_match:
                            scene = scene_match.group(1)
                    
                    # Add scene header
                    doc.add_heading(f"Act {act}, Scene {scene}", level=1)
                    
                    # Process markdown content
                    lines = md_content.split('\n')
                    
                    # Skip act and scene headers
                    content_lines = []
                    for line in lines:
                        if re.match(r'^ACT\s+[IVX\d]+', line, re.IGNORECASE):
                            continue
                        if re.match(r'^SCENE\s+[IVX\d]+', line, re.IGNORECASE):
                            continue
                        content_lines.append(line)
                    
                    # Add content
                    for line in content_lines:
                        line = line.strip()
                        if not line:
                            continue
                            
                        # Check if it's a stage direction [...]
                        if line.startswith('[') and line.endswith(']'):
                            p = doc.add_paragraph()
                            run = p.add_run(line)
                            run.italic = True
                            continue
                            
                        # Check if it's a character name (all caps)
                        if line.isupper() and len(line.split()) <= 3:
                            p = doc.add_paragraph()
                            run = p.add_run(line)
                            run.bold = True
                            p.space_after = Pt(0)  # No space after character name
                            continue
                            
                        # Regular dialogue
                        p = doc.add_paragraph(line)
                        p.paragraph_format.left_indent = Pt(36)  # Indent dialogue
                
                # Add page break after each scene
                doc.add_page_break()
                
            except Exception as e:
                print(f"Error processing scene {scene_path}: {str(e)}")
                # Continue with next scene
        
        # Save the document
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        return output_path

    def export_play_from_markdown(self, md_path: str, output_path: str) -> str:
        """
        Export a complete play from a combined markdown file to DOCX format.
        
        Args:
            md_path: Path to the combined markdown file
            output_path: Path for the output file
            
        Returns:
            Path to the created file
        """
        if not self.docx_available:
            raise ImportError("python-docx is required for DOCX export")
        
        # Load markdown content
        try:
            with open(md_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
        except Exception as e:
            raise ValueError(f"Error loading markdown file: {str(e)}")
        
        # Extract play title (if present)
        title_match = re.search(r'^#\s+(.+)$', md_content, re.MULTILINE)
        title = title_match.group(1) if title_match else "Play"
        
        # Create DOCX
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Add title page
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add break for title page
        doc.add_page_break()
        
        # Process content by sections (Act/Scene blocks)
        current_act = None
        current_scene = None
        dialogue_buffer = []
        
        # Split content into lines and process
        lines = md_content.split('\n')
        
        for i, line in enumerate(lines):
            # Skip the title line we already processed
            if i == 0 and title_match and title_match.group(0) == line:
                continue
                
            # Check for Act header
            act_match = re.match(r'^(?:#+\s+)?ACT\s+([IVX\d]+)', line, re.IGNORECASE)
            if act_match:
                # When we find a new act, add the previous scene's content
                if dialogue_buffer:
                    self._add_dialogue_to_doc(doc, dialogue_buffer)
                    dialogue_buffer = []
                
                current_act = act_match.group(1)
                doc.add_heading(f"Act {current_act}", level=1)
                continue
                
            # Check for Scene header
            scene_match = re.match(r'^(?:#+\s+)?SCENE\s+([IVX\d]+)', line, re.IGNORECASE)
            if scene_match:
                # When we find a new scene, add the previous scene's content
                if dialogue_buffer:
                    self._add_dialogue_to_doc(doc, dialogue_buffer)
                    dialogue_buffer = []
                
                current_scene = scene_match.group(1)
                doc.add_heading(f"Scene {current_scene}", level=2)
                continue
                
            # Add other content to the dialogue buffer
            line = line.strip()
            if line:
                dialogue_buffer.append(line)
        
        # Add any remaining content
        if dialogue_buffer:
            self._add_dialogue_to_doc(doc, dialogue_buffer)
        
        # Save the document
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        return output_path

    def _add_dialogue_to_doc(self, doc, lines):
        """Helper method to add dialogue lines to document with proper formatting."""
        from docx.shared import Pt
        
        for line in lines:
            # Check if it's a stage direction [...]
            if line.startswith('[') and line.endswith(']'):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.italic = True
                continue
                
            # Check if it's a character name (all caps)
            if line.isupper() and len(line.split()) <= 3:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                p.space_after = Pt(0)  # No space after character name
                continue
                
            # Regular dialogue
            p = doc.add_paragraph(line)
            p.paragraph_format.left_indent = Pt(36)  # Indent dialogue