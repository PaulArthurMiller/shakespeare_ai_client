"""
Export Manager for Shakespeare AI.

This module handles exporting plays and scenes to various formats,
as well as combining multiple scenes into a complete play.
"""
import os
import shutil
import re
from typing import Dict, List, Any, Optional, Tuple, Union

from modules.ui.file_helper import (
    load_text_from_file,
    save_text_to_file,
    ensure_directory,
    extract_act_scene_from_filename
)


class ExportManager:
    """
    Handles export operations for the Shakespeare AI playwright.
    """
    
    def __init__(self, logger=None):
        """
        Initialize the export manager.
        
        Args:
            logger: Optional logger instance
        """
        self.logger = logger
        
        # Check if docx support is available
        try:
            import docx
            self.docx_available = True
        except ImportError:
            self.docx_available = False
            self._log("python-docx not available. Install with: pip install python-docx")
    
    def _log(self, message: str, level: str = "info") -> None:
        """Log a message using the provided logger if available."""
        if self.logger:
            if hasattr(self.logger, "_log"):
                self.logger._log(message, level)
            elif hasattr(self.logger, level):
                getattr(self.logger, level)(message)
        else:
            print(f"[{level.upper()}] {message}")
    
    def _act_to_int(self, act: str) -> int:
        """
        Convert act identifier to integer for sorting.
        
        Args:
            act: Act identifier
            
        Returns:
            Integer value
        """
        # Check if it's a Roman numeral
        if re.match(r'^[IVXLCDM]+$', act.upper()):
            romans = {'I': 1, 'V': 5, 'X': 10, 'L': 50, 'C': 100, 'D': 500, 'M': 1000}
            result = 0
            prev = 0
            for c in act.upper():
                if c in romans:
                    current = romans[c]
                    if current > prev:
                        result += current - 2 * prev
                    else:
                        result += current
                    prev = current
            return result
        
        # Try to convert to integer directly
        try:
            return int(act)
        except ValueError:
            # If all else fails, return a large number to sort it last
            return 9999
    
    def _scene_to_int(self, scene: str) -> int:
        """Alias for _act_to_int for scene identifiers."""
        return self._act_to_int(scene)
    
    def combine_scenes(self, base_output_dir: str, 
                      output_filename: Optional[str] = None) -> Tuple[bool, str]:
        """
        Combine all generated scenes into a single play file.
        
        Args:
            base_output_dir: Base directory for input/output
            output_filename: Optional output filename
            
        Returns:
            Tuple of (success, error_message or output_path)
        """
        try:
            self._log("Combining scenes into a single play file...")
            
            # Set default output filename if not provided
            if output_filename is None:
                output_filename = "modern_play_combined.md"
            
            # Full output path
            output_path = os.path.join(base_output_dir, output_filename)
            
            # Get scene files
            scenes_dir = os.path.join(base_output_dir, "generated_scenes_claude2")
            if not os.path.exists(scenes_dir):
                scenes_dir = os.path.join(base_output_dir, "generated_scenes")
            
            if not os.path.exists(scenes_dir):
                return False, f"Scenes directory not found at {scenes_dir}"
            
            # Get scene files and sort them
            scene_files = []
            for filename in os.listdir(scenes_dir):
                if filename.endswith(".md"):
                    filepath = os.path.join(scenes_dir, filename)
                    act, scene = extract_act_scene_from_filename(filename)
                    scene_files.append((filepath, filename, act, scene))
            
            if not scene_files:
                return False, f"No scene files found in {scenes_dir}"
            
            # Sort scene files by act and scene
            scene_files.sort(key=lambda x: (self._act_to_int(x[2]), self._scene_to_int(x[3])))
            
            # Combine the files
            combined_text = ""
            for filepath, _, _, _ in scene_files:
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        scene_text = f.read().strip()
                    combined_text += scene_text + "\n\n"
                except Exception as e:
                    self._log(f"Error reading scene file {filepath}: {e}", "error")
            
            # Save the combined play
            if not save_text_to_file(combined_text, output_path):
                return False, f"Error saving combined play to {output_path}"
            
            self._log(f"Combined play saved to {output_path}")
            return True, output_path
        except Exception as e:
            error_msg = f"Error combining scenes: {str(e)}"
            self._log(error_msg, "error")
            return False, error_msg
    
    def combine_scenes_in_project(self, project_id: str, 
                                output_filename: str) -> Tuple[bool, str]:
        """
        Combine all scenes in a project into a single play file.
        
        Args:
            project_id: Project identifier
            output_filename: Output filename
            
        Returns:
            Tuple of (success, error_message or output_path)
        """
        try:
            self._log(f"Combining scenes for project {project_id}")
            
            # Project folder
            project_folder = os.path.join("data/play_projects", project_id)
            project_scenes_dir = os.path.join(project_folder, "scenes")
            
            if not os.path.exists(project_scenes_dir):
                return False, f"Scenes directory not found at {project_scenes_dir}"
            
            # Get scene files and sort them
            scene_files = []
            for filename in os.listdir(project_scenes_dir):
                if filename.endswith(".md"):
                    filepath = os.path.join(project_scenes_dir, filename)
                    act, scene = extract_act_scene_from_filename(filename)
                    scene_files.append((filepath, filename, act, scene))
            
            if not scene_files:
                return False, f"No scene files found in {project_scenes_dir}"
            
            # Sort scene files by act and scene
            scene_files.sort(key=lambda x: (self._act_to_int(x[2]), self._scene_to_int(x[3])))
            
            # Get project data for title
            from modules.ui.playwright.project_manager import ProjectManager
            project_manager = ProjectManager(logger=self.logger)
            project_data = project_manager.get_project_data(project_id)
            
            # Use project title as header if available
            title = project_data.get("title", "Play") if project_data else "Play"
            combined_text = f"# {title}\n\n"
            
            # Combine the files
            for filepath, _, _, _ in scene_files:
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        scene_text = f.read().strip()
                    combined_text += scene_text + "\n\n"
                except Exception as e:
                    self._log(f"Error reading scene file {filepath}: {e}", "error")
            
            # Full output path
            output_path = os.path.join(project_folder, output_filename)
            
            # Save the combined play
            if not save_text_to_file(combined_text, output_path):
                return False, f"Error saving combined play to {output_path}"
            
            self._log(f"Combined play saved to {output_path}")
            return True, output_path
        except Exception as e:
            error_msg = f"Error combining scenes: {str(e)}"
            self._log(error_msg, "error")
            return False, error_msg
    
    def save_scene_to_file(self, project_id: str, act: str, scene: str, 
                        output_format: str = "docx") -> Tuple[bool, str]:
        """
        Save a specific scene to a file in the desired format.
        
        Args:
            project_id: Project identifier
            act: Act identifier
            scene: Scene identifier
            output_format: Output format ("docx" or "md")
            
        Returns:
            Tuple of (success, output_path)
        """
        project_folder = os.path.join("data/play_projects", project_id)
        scene_md_path = os.path.join(project_folder, "scenes", 
                                f"act_{act.lower()}_scene_{scene.lower()}.md")
        
        # Check if scene exists
        if not os.path.exists(scene_md_path):
            return False, f"Scene file not found: {scene_md_path}"
        
        # Create output directory
        output_dir = os.path.join(project_folder, "exports")
        ensure_directory(output_dir)
        
        # Save logs with the export
        self.save_logs_with_export(project_id, output_dir)
        
        if output_format.lower() == "md":
            # For markdown, just copy the file
            output_path = os.path.join(output_dir, f"act_{act.lower()}_scene_{scene.lower()}.md")
            try:
                shutil.copy2(scene_md_path, output_path)
                return True, output_path
            except Exception as e:
                return False, f"Error copying file: {str(e)}"
        
        elif output_format.lower() == "docx":
            # For docx, use an exporter (assume it's installed or import if possible)
            if self.docx_available:
                try:
                    from docx import Document
                    from docx.shared import Pt, RGBColor
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    
                    # Read the markdown
                    content = load_text_from_file(scene_md_path)
                    if not content:
                        return False, "Could not read scene file"
                    
                    # Create the document
                    doc = Document()
                    
                    # Title
                    title = doc.add_heading(f"Act {act}, Scene {scene}", level=1)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Process the markdown line by line
                    lines = content.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        
                        # Skip act and scene headers that are already in the title
                        if re.match(r'^ACT\s+[IVX\d]+', line, re.IGNORECASE) or \
                        re.match(r'^SCENE\s+[IVX\d]+', line, re.IGNORECASE):
                            continue
                            
                        # Check if it's a stage direction [...]
                        if line.startswith('[') and line.endswith(']'):
                            p = doc.add_paragraph()
                            run = p.add_run(line)
                            run.italic = True
                            continue
                            
                        # Check if it's a character name (all caps)
                        if line.isupper() and len(line.split()) <= 3:
                            p = doc.add_paragraph()
                            run = p.add_run(line)
                            run.bold = True
                            p.paragraph_format.space_after = Pt(0)
                            continue
                            
                        # Regular dialogue
                        p = doc.add_paragraph(line)
                        p.paragraph_format.left_indent = Pt(36)  # Indent dialogue
                    
                    # Save the document
                    output_path = os.path.join(output_dir, f"act_{act.lower()}_scene_{scene.lower()}.docx")
                    doc.save(output_path)
                    
                    return True, output_path
                except Exception as e:
                    return False, f"Error creating DOCX: {str(e)}"
            else:
                return False, "python-docx not available. Install with: pip install python-docx"
        else:
            return False, f"Unsupported output format: {output_format}"

    def save_full_play_to_file(self, project_id: str, 
                            output_format: str = "docx") -> Tuple[bool, str]:
        """
        Save all scenes as a full play file in the desired format.
        
        Args:
            project_id: Project identifier
            output_format: Output format ("docx" or "md")
            
        Returns:
            Tuple of (success, output_path)
        """
        # Get project data
        from modules.ui.playwright.project_manager import ProjectManager
        project_manager = ProjectManager(logger=self.logger)
        project_data = project_manager.get_project_data(project_id)
        
        if not project_data:
            return False, f"Project not found: {project_id}"
        
        project_folder = os.path.join("data/play_projects", project_id)
        project_scenes_dir = os.path.join(project_folder, "scenes")
        
        # Create output directory
        output_dir = os.path.join(project_folder, "exports")
        ensure_directory(output_dir)
        
        # Get title for filename - ensure it's a string and sanitize for filename use
        play_title = project_data.get("title", "play")
        if not isinstance(play_title, str):
            play_title = "play"  # Default if title is not a string
        safe_title = str(play_title).replace(" ", "_").lower()
        
        # Save logs with the export
        self.save_logs_with_export(project_id, output_dir)
        
        if output_format.lower() == "md":
            # First combine scenes into a single markdown file
            success, result = self.combine_scenes_in_project(
                project_id=project_id,
                output_filename=f"{safe_title}_full.md"
            )
            
            if not success:
                return False, result
                
            # Copy to exports directory
            source_path = result
            output_path = os.path.join(output_dir, f"{safe_title}_full.md")
            
            try:
                shutil.copy2(source_path, output_path)
                return True, output_path
            except Exception as e:
                return False, f"Error copying file: {str(e)}"
            
        elif output_format.lower() == "docx":
            # For docx, we need to create a new document
            if not self.docx_available:
                return False, "python-docx not available. Install with: pip install python-docx"
                
            try:
                from docx import Document
                from docx.shared import Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                # Get all scene files
                scene_files = []
                for filename in os.listdir(project_scenes_dir):
                    if filename.endswith(".md"):
                        filepath = os.path.join(project_scenes_dir, filename)
                        act, scene = extract_act_scene_from_filename(filename)
                        scene_files.append((filepath, filename, act, scene))
                
                # Sort scene files
                scene_files.sort(key=lambda x: (self._act_to_int(x[2]), self._scene_to_int(x[3])))
                
                if not scene_files:
                    return False, "No scene files found"
                
                # Create document
                doc = Document()
                
                # Add title - Get title as a string
                title_text = str(project_data.get("title", "Play"))
                title = doc.add_heading(title_text, level=0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add page break after title
                doc.add_page_break()
                
                # Process each scene
                current_act = None
                
                for filepath, _, act, scene in scene_files:
                    # Add act header if it's a new act
                    if act != current_act:
                        doc.add_heading(f"Act {act}", level=1)
                        current_act = act
                    
                    # Add scene header
                    doc.add_heading(f"Scene {scene}", level=2)
                    
                    # Read scene content
                    content = load_text_from_file(filepath)
                    if not content:
                        self._log(f"Could not read scene file: {filepath}", "warning")
                        continue
                    
                    # Process the scene content line by line
                    lines = content.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        
                        # Skip act and scene headers that are already added
                        if re.match(r'^ACT\s+[IVX\d]+', line, re.IGNORECASE) or \
                        re.match(r'^SCENE\s+[IVX\d]+', line, re.IGNORECASE):
                            continue
                            
                        # Check if it's a stage direction [...]
                        if line.startswith('[') and line.endswith(']'):
                            p = doc.add_paragraph()
                            run = p.add_run(line)
                            run.italic = True
                            continue
                            
                        # Check if it's a character name (all caps)
                        if line.isupper() and len(line.split()) <= 3:
                            p = doc.add_paragraph()
                            run = p.add_run(line)
                            run.bold = True
                            p.paragraph_format.space_after = Pt(0)
                            continue
                            
                        # Regular dialogue
                        p = doc.add_paragraph(line)
                        p.paragraph_format.left_indent = Pt(36)  # Indent dialogue
                    
                    # Add page break after each scene
                    doc.add_page_break()
                
                # Save the document - USE THE SAFE TITLE HERE
                output_path = os.path.join(output_dir, f"{safe_title}_full.docx")
                doc.save(output_path)
                
                return True, output_path
                
            except Exception as e:
                return False, f"Error creating DOCX: {str(e)}"
        else:
            return False, f"Unsupported output format: {output_format}"
        
    def save_logs_with_export(self, project_id: str, output_dir: str) -> None:
        """
        Copy the project logs to the export directory.
        
        Args:
            project_id: Project identifier
            output_dir: Export output directory
        """
        try:
            # Get the project logs directory
            project_logs_dir = os.path.join("data/play_projects", project_id, "logs")
            if not os.path.exists(project_logs_dir):
                self._log(f"No logs directory found for project: {project_id}")
                return
                
            # Create logs directory in export directory
            export_logs_dir = os.path.join(output_dir, "logs")
            ensure_directory(export_logs_dir)
            
            # Copy all log files
            for log_file in os.listdir(project_logs_dir):
                if log_file.endswith(".log"):
                    src = os.path.join(project_logs_dir, log_file)
                    dst = os.path.join(export_logs_dir, log_file)
                    shutil.copy2(src, dst)
            
            self._log(f"Copied project logs to export directory: {export_logs_dir}")
        except Exception as e:
            self._log(f"Error copying logs: {e}", "error")